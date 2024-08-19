import tkinter as tk
from tkinter import scrolledtext, messagebox
from tkinter import ttk
from tkinter import filedialog
from docx import Document
from openpyxl import Workbook
from openpyxl.styles import Alignment
import re

# Paths to the DOCX files
SRD_MFD_PATH = r'C:\Users\lsrinivasan\Desktop\Trace_documents\SRD_MFD.docx'
SDD_MFD_PATH = r'C:\Users\lsrinivasan\Desktop\Trace_documents\SDD_MFD.docx'
TC_PATH = r'C:\Users\lsrinivasan\Desktop\Trace_documents\TC_MFD_GROUND_SPEED.docx'
TEC_SPEC_PATH = r'C:\Users\lsrinivasan\Desktop\Trace_documents\Tech Spec MFD.docx'


def read_requirement_from_srd(requirement_number):
    document = Document(SRD_MFD_PATH)
    full_text = ""
    capture = False
    for para in document.paragraphs:
        if requirement_number in para.text:
            capture = True
        if capture:
            full_text += para.text + "\n"
            if not para.text.strip() and full_text.strip():  # Stop capturing after a blank line if text is captured
                break
    return full_text.strip() if full_text else "Requirement not found."


def read_trace_from_sdd(requirement_number):
    document = Document(SDD_MFD_PATH)
    full_text = ""
    current_trace = ""
    capture = False

    for para in document.paragraphs:
        if "Req ID:" in para.text:
            current_trace = ""
            capture = False
        current_trace += para.text + "\n"
        if requirement_number in para.text:
            capture = True
            full_text += current_trace + "\n"
            current_trace = ""

    return full_text.strip() if full_text else "Trace not found."


def read_trace_from_tec_spec(trace_numbers):
    document = Document(TEC_SPEC_PATH)
    trace_texts = []
    current_trace = ""
    capture = False

    # Compile regex patterns to match exact trace numbers
    trace_number_patterns = [re.compile(r'^\s*' + re.escape(trace_number) + r'(\s|$)', re.IGNORECASE) for trace_number
                             in trace_numbers]

    for para in document.paragraphs:
        # Start capturing if any trace number pattern matches the start of the paragraph text
        if any(pattern.match(para.text) for pattern in trace_number_patterns):
            if current_trace.strip():
                trace_texts.append(current_trace.strip())
            current_trace = para.text + "\n"
            capture = True
        elif capture:
            # Stop capturing when reaching "REQ Verification: Test"
            if "REQ Verification: Test" in para.text:
                current_trace += para.text + "\n"
                trace_texts.append(current_trace.strip())
                current_trace = ""
                capture = False
            else:
                current_trace += para.text + "\n"

    # Append any remaining captured text
    if current_trace.strip():
        trace_texts.append(current_trace.strip())

    return "\n\n".join(trace_texts) if trace_texts else "Trace not found."


def extract_tables_from_docx(docx_path):
    document = Document(docx_path)
    tables = []
    for table in document.tables:
        table_data = []
        for row in table.rows:
            row_data = [cell.text.strip() for cell in row.cells]
            table_data.append(row_data)
        tables.append(table_data)
    return tables


def find_tables_by_requirement(tables, requirement_number):
    matching_tables = []
    for table in tables:
        for row in table:
            if row and len(row) > 0 and requirement_number in row[-1]:  # Check if requirement number is in the last column
                matching_tables.append(table)
                break
    return matching_tables


def search_requirement():
    requirement_number = requirement_combobox.get()
    srd_text = read_requirement_from_srd(requirement_number)
    sdd_text = read_trace_from_sdd(requirement_number)

    trace_numbers = []
    if "[Trace:" in srd_text:
        trace_section = srd_text.split("[Trace:")[1].split("]")[0]
        trace_numbers = [trace.strip() for trace in trace_section.split(",")]

    tec_spec_text = read_trace_from_tec_spec(trace_numbers)

    # Display SRD, SDD, and Tec_Spec text
    display_formatted_text(srd_display,  srd_text)
    display_formatted_text(sdd_display,  sdd_text)
    display_formatted_text(tec_spec_display,  tec_spec_text)

    # Find and display table data
    tables = extract_tables_from_docx(TC_PATH)
    matched_tables = find_tables_by_requirement(tables, requirement_number)

    if matched_tables:
        table_text = "\n\n".join(
            f"Table Data related to {requirement_number}:\n" + "-" * 80 + "\n" + format_table_as_text(table) for table in matched_tables)
    else:
        table_text = "Requirement not found."

    display_formatted_text(table_display, table_text)


def format_table_as_text(table):
    text = ""
    for row in table:
        text += " | ".join(row) + "\n"
    return text.strip()


def display_formatted_text(display_widget, text):
    display_widget.configure(state='normal')
    display_widget.delete(1.0, tk.END)
    display_widget.insert(tk.END, text)
    display_widget.configure(state='disabled')


def save_to_excel():
    requirement_number = requirement_combobox.get()
    srd_text = read_requirement_from_srd(requirement_number)
    sdd_text = read_trace_from_sdd(requirement_number)
    tables = extract_tables_from_docx(TC_PATH)
    matched_tables = find_tables_by_requirement(tables, requirement_number)

    # Prompt user to select save location
    file_path = filedialog.asksaveasfilename(defaultextension=".xlsx",
                                             filetypes=[("Excel files", "*.xlsx"), ("All files", "*.*")])
    if not file_path:
        return

    # Create a new Excel workbook and add data
    workbook = Workbook()
    sheet = workbook.active
    sheet.title = "Requirement Data"

    # Add SRD data
    sheet.append(["SRD_MFD Requirement"])
    srd_text_lines = srd_text.split('\n')
    for line in srd_text_lines:
        sheet.append([line])
    sheet.append([])  # Add an empty row for spacing

    # Add SDD data
    sheet.append(["SDD_MFD Corresponding Requirements"])
    sdd_text_lines = sdd_text.split('\n')
    for line in sdd_text_lines:
        sheet.append([line])
    sheet.append([])  # Add an empty row for spacing

    # Add Tec_Spec data
    sheet.append(["Tec_Spec Corresponding Requirements"])
    tec_spec_text_lines = tec_spec_display.get(1.0, tk.END).split('\n')
    for line in tec_spec_text_lines:
        sheet.append([line])
    sheet.append([])  # Add an empty row for spacing

    # Add table data
    for i, table in enumerate(matched_tables):
        sheet.append([f'Table {i + 1} Data'])
        for row in table:
            sheet.append(row)
        sheet.append([])  # Add an empty row for spacing

    # Adjust column width and alignment
    for col in sheet.columns:
        max_length = 0
        column = col[0].column_letter  # Get the column name
        for cell in col:
            try:
                if len(str(cell.value)) > max_length:
                    max_length = len(cell.value)
            except:
                pass
        adjusted_width = (max_length + 2)
        sheet.column_dimensions[column].width = adjusted_width

    for row in sheet.iter_rows():
        for cell in row:
            cell.alignment = Alignment(wrapText=True, vertical='top')

    # Save the workbook
    workbook.save(file_path)
    messagebox.showinfo("Save Successful", f"Data saved to '{file_path}'")


def previous_requirement():
    current_index = requirement_combobox.current()
    if current_index > 0:
        requirement_combobox.current(current_index - 1)
        search_requirement()


def next_requirement():
    current_index = requirement_combobox.current()
    if current_index < len(requirement_numbers) - 1:
        requirement_combobox.current(current_index + 1)
        search_requirement()


# Create the main window
window = tk.Tk()
window.title("Document Search Tool")

# Set the window size
window.geometry('1200x1000')

# Create the search bar and button
tk.Label(window, text="Select Requirement Number:").grid(row=0, column=0, padx=10, pady=5)

# Generate requirement numbers
requirement_numbers = [f'MFD_HLR_{i:03}' for i in range(1, 465)]  # Adjust range based on your data

# Create a combobox for requirement numbers
requirement_combobox = ttk.Combobox(window, values=requirement_numbers)
requirement_combobox.grid(row=0, column=1, padx=10, pady=5)
requirement_combobox.current(0)  # Set default selection to the first item

search_button = tk.Button(window, text="Search", command=search_requirement)
search_button.grid(row=0, column=2, padx=10, pady=5)

save_button = tk.Button(window, text="Save", command=save_to_excel)
save_button.grid(row=0, column=3, padx=10, pady=5)

previous_button = tk.Button(window, text="Previous", command=previous_requirement)
previous_button.grid(row=0, column=4, padx=10, pady=5)

next_button = tk.Button(window, text="Next", command=next_requirement)
next_button.grid(row=0, column=5, padx=10, pady=5)

# Create the display areas
tk.Label(window, text="SRD_MFD Requirement Text:").grid(row=1, column=0, padx=10, pady=5)
srd_display = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=150, height=8, font=("Times New Roman", 11))
srd_display.grid(row=2, column=0, columnspan=6, padx=10, pady=5)

tk.Label(window, text="SDD_MFD Corresponding Requirements:").grid(row=3, column=0, padx=10, pady=5)
sdd_display = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=150, height=8, font=("Times New Roman", 11))
sdd_display.grid(row=4, column=0, columnspan=6, padx=10, pady=5)

tk.Label(window, text="Tec_Spec Corresponding Requirements:").grid(row=5, column=0, padx=10, pady=5)
tec_spec_display = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=150, height=8, font=("Times New Roman", 11))
tec_spec_display.grid(row=6, column=0, columnspan=6, padx=10, pady=5)

tk.Label(window, text="Test Case Table Data:").grid(row=7, column=0, padx=10, pady=5)
table_display = scrolledtext.ScrolledText(window, wrap=tk.WORD, width=150, height=8, font=("Times New Roman", 11))
table_display.grid(row=8, column=0, columnspan=6, padx=10, pady=5)

# Start the GUI event loop
window.mainloop()
